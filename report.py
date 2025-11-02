"""
generate_tripnbook_report_v2.py

Improved TripNBook report generator:
 - strict formatting (margins, font, line spacing)
 - black titles/headings only
 - Table of Contents field (open in Word and update TOC)
 - footer with page numbers
 - expanded content + sample Ansible/Nagios/AI snippets
 - diagram placeholders or embed by setting diagram_paths

Usage:
    pip install python-docx
    python generate_tripnbook_report_v2.py
Output:
    TripNBook_Report_v2.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime
import os

# === User settings ===
OUTPUT_FILENAME = "TripNBook_Report_v2.docx"

diagram_paths = {
    "ER_DIAGRAM": None,      # e.g., "images/erd.png"
    "DFD_LEVEL_0": None,
    "DFD_LEVEL_1": None,
}

PROJECT_TITLE = "TripNBook — AI-Powered Flight Booking Website"
GROUP_MEMBERS = ["Vinod Ravi (Lead Developer)"]
SUPERVISOR = "Project Supervisor: ____________________"
UNIVERSITY = "Lovely Professional University (LPU)"
COURSE = "INT 333 — Project Report"
SUBMISSION_DATE = datetime.date.today().strftime("%B %d, %Y")

# === Helpers ===
def set_default_font(doc, name='Times New Roman', size=12):
    style = doc.styles['Normal']
    font = style.font
    font.name = name
    font.size = Pt(size)
    font._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def set_heading_styles():
    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0,0,0)
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0,0,0)
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0,0,0)
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_centered(document, text, size=16, bold=False, spacing_after=12):
    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = RGBColor(0,0,0)
    p.space_after = Pt(spacing_after)
    return p

def add_heading(text, level=1):
    # use document.add_heading so Word recognizes it for TOC
    h = doc.add_heading(text if level==1 else text, level=level)
    # force color & font on the run(s) inside heading
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(16 if level==1 else (14 if level==2 else 12))
        run.font.bold = True
        run.font.color.rgb = RGBColor(0,0,0)
    return h

def add_para(text, first_line_indent=False, style_name=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    return p

def insert_image_or_placeholder(image_path, caption):
    if image_path and os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_picture(image_path, width=Inches(5.5))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cap.runs[0].italic = True
    else:
        add_para(f"[Insert {caption} here — recommended export from eraser.io (1600×900).]", first_line_indent=False)

def add_toc_field():
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)
    add_para("Right-click the table of contents and choose 'Update Field' or press F9 in Word to generate page numbers.", first_line_indent=False)

def add_page_number_footer():
    section = doc.sections[-1]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # create field for page number
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), "PAGE")
    p._p.append(fldSimple)

# === Build document ===
doc = Document()
set_default_font(doc)
set_heading_styles()

# margins
section = doc.sections[0]
section.top_margin = Inches(1)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)

# Cover page
add_centered(doc, UNIVERSITY, size=18, bold=True)
add_centered(doc, COURSE, size=14, bold=True)
doc.add_paragraph()
add_centered(doc, PROJECT_TITLE, size=20, bold=True)
add_centered(doc, "Project Report", size=14)
doc.add_paragraph()
for m in GROUP_MEMBERS:
    add_centered(doc, m, size=12)
add_centered(doc, SUPERVISOR, size=12)
add_centered(doc, f"Submission Date: {SUBMISSION_DATE}", size=12)
add_para("Note: Cover styling (page background color, stamp) can be set in Word to match university specification.")
doc.add_page_break()

# Preliminary pages placeholder
add_para("Declaration ............................................................................... (i)")
add_para("Certificate ............................................................................. (ii)")
add_para("Acknowledgement ..................................................................... (iii)")
doc.add_page_break()

# TOC
add_para("TABLE OF CONTENTS", first_line_indent=False)
add_toc_field()
doc.add_page_break()

# --- CHAPTER 1: INTRODUCTION ---
add_heading("1. INTRODUCTION", level=1)
add_para("TripNBook is an AI-first flight-only booking platform designed for speed, resilience, and a frictionless user experience. The application focuses on core flight flows (search, seat selection, booking) while embedding an intelligent assistant to interpret fuzzy queries, correct typos, and recommend alternatives when exact matches are not found.", first_line_indent=True)

add_heading("1.1 Background", level=2)
add_para("Air travel search requires high-quality parsing of origin/destination and date constraints. Users often mistype city names, abbreviations, or airport codes; a system that corrects and suggests reduces abandonment and improves conversion. TripNBook's AI assistant bridges this UX gap.", first_line_indent=True)

add_heading("1.2 Objectives", level=2)
add_para("Primary objectives:\n• Fast and tolerant flight search (typo recovery + suggestions).\n• Clear passenger and seat booking flows with realistic aircraft seatmaps.\n• Repeatable deployments via Ansible and operational monitoring using Nagios.\n• Maintainable, tested codebase ready for campus/institutional evaluation.", first_line_indent=True)

doc.add_page_break()

# === CHAPTER 2: PROBLEM STATEMENT & RATIONALE ===
add_heading("2. PROBLEM STATEMENT & RATIONALE", level=1)
add_para("Problem: Users commonly abandon searches when the system returns 'no results' for a misspelled or partially-entered origin/destination. Further, ad-hoc deployments without monitoring increase time-to-recovery on incidents. TripNBook's purpose is to address these UX and operational shortcomings in a compact, teachable MVP.", first_line_indent=True)

add_para("Rationale: Focusing on a flight-only product allows deeper polishing of search quality and operational reliability. This makes the project feasible within academic timelines while delivering a production-informed architecture.", first_line_indent=True)

doc.add_page_break()

# === CHAPTER 3: LITERATURE / EXISTING SYSTEMS & DIAGRAMS ===
add_heading("3. EXISTING SYSTEMS & DIAGRAMS", level=1)
add_para("Major OTAs provide comprehensive functionality but come with heavy infra and complexity. TripNBook intentionally chooses a minimal stack to demonstrate key engineering principles.", first_line_indent=True)
add_para("Diagrams (ERD, DFD) are required by the university. Insert the exported images below or use the placeholders.", first_line_indent=True)
insert_image_or_placeholder(diagram_paths.get("ER_DIAGRAM"), "ER Diagram")
insert_image_or_placeholder(diagram_paths.get("DFD_LEVEL_0"), "DFD Level 0")
insert_image_or_placeholder(diagram_paths.get("DFD_LEVEL_1"), "DFD Level 1")
add_para("Diagrams should include entities: Flight, AircraftModel, Seat, Booking, Passenger, User, and logs/monitoring entities for Nagios alerts.", first_line_indent=True)

doc.add_page_break()

# === CHAPTER 4: DETAILED ANALYSIS & DATA MODEL ===
add_heading("4. DETAILED ANALYSIS & DATA MODEL", level=1)
add_para("Data model overview: The core collection is 'flights' with schema fields like flight_no, origin, destination, depart_time, arrival_time, aircraft_model, seat_map (json), fares. Booking documents reference flights and contain passenger arrays, PNR, payment_status, and timestamps.", first_line_indent=True)
add_para("Security & PII: Only store minimal passenger PII required for booking (name, DOB, passport), enforced via environment-variable-controlled encryption at rest (e.g., MongoDB Atlas encryption) and transport-level TLS.", first_line_indent=True)

doc.add_page_break()

# === CHAPTER 5: SOFTWARE REQUIREMENTS & NON-FUNCTIONALS ===
add_heading("5. SOFTWARE REQUIREMENTS & NON-FUNCTIONALS", level=1)
add_para("Functional Requirements (expanded):\n- FR1: Typo-tolerant search that returns ranked candidate airports/cities.\n- FR2: Realistic seatmap rendering driven by aircraftLayouts config.\n- FR3: Complete booking lifecycle with status updates and email confirmations.\n- FR4: Admin view to reconcile bookings and view Nagios alerts.", first_line_indent=True)

add_para("Non-functional Requirements (expanded):\n- NFR1: Search latency < 2s under normal load.\n- NFR2: 99% availability for the web UI with Nagios alerts configured.\n- NFR3: Basic accessibility and responsive design for mobile and desktop.", first_line_indent=True)

doc.add_page_break()

# === CHAPTER 6: SYSTEM DESIGN & AI MODULE ===
add_heading("6. SYSTEM DESIGN & AI MODULE", level=1)
add_para("Architecture summary: React frontend (CRA) <-> Express/Node backend <-> MongoDB Atlas. Optional containerization with Docker; Ansible playbooks to provision VMs/containers; Nagios monitors hosts and service endpoints.", first_line_indent=True)

add_heading("6.1 AI Assistant (fuzzy matching)", level=2)
add_para("Design: The AI assistant normalizes input, applies fuzzy string matching (Levenshtein/edit distance), weighted heuristics (population, popularity of airports), and returns top-K candidates. It can be a microservice (Flask/FastAPI) or integrated in backend.", first_line_indent=True)

add_para("Sample Python fuzzy-match (conceptual):", first_line_indent=False)
add_para(
"""# conceptual example (not a full dependency file)
from difflib import get_close_matches
def fuzzy_airport_search(query, airports, n=5):
    q = query.strip().lower()
    names = [a['name'].lower() for a in airports]
    matches = get_close_matches(q, names, n=n, cutoff=0.6)
    # return original airport records for matches
    return [a for a in airports if a['name'].lower() in matches]
""", first_line_indent=True)

doc.add_page_break()

# === CHAPTER 7: IMPLEMENTATION, DEVOPS & MONITORING ===
add_heading("7. IMPLEMENTATION, DEVOPS & MONITORING", level=1)
add_para("Deployment approach: Ansible playbooks configure base VMs, install Node, pull repo, build frontend, run backend as a systemd service or Docker container. The playbooks are idempotent for repeatable deployments.", first_line_indent=True)

add_heading("7.1 Sample Ansible playbook snippet", level=2)
add_para("This is a minimal deploy playbook that illustrates roles and tasks. Put it in playbooks/deploy.yml and run via `ansible-playbook -i inventory playbooks/deploy.yml`.", first_line_indent=True)
add_para(
"""---
- hosts: webservers
  become: yes
  vars:
    repo_url: 'https://github.com/youruser/TripNBook.git'
  tasks:
    - name: Ensure node is installed (example)
      apt:
        name: nodejs
        state: present
    - name: Clone repository
      git:
        repo: "{{ repo_url }}"
        dest: /opt/tripnbook
        version: main
    - name: Install backend dependencies
      command: npm install
      args:
        chdir: /opt/tripnbook/backend
    - name: Build frontend
      command: npm run build
      args:
        chdir: /opt/tripnbook/frontend
""", first_line_indent=True)

add_heading("7.2 Nagios monitoring examples", level=2)
add_para("Nagios checks to configure:\n• check_http for /health endpoints\n• NRPE check for backend process and queue length\n• Custom script (check_booking_queue.sh) returning CRITICAL when > X pending", first_line_indent=True)
add_para("Example Nagios command definition (commands.cfg):", first_line_indent=False)
add_para(
"""define command{
  command_name check_tripnbook_health
  command_line /usr/lib/nagios/plugins/check_http -I $HOSTADDRESS$ -u /health -w 5 -c 10
}""", first_line_indent=True)

doc.add_page_break()

# === CHAPTER 8: TESTING, QA & MAINTENANCE ===
add_heading("8. TESTING, QA & MAINTENANCE", level=1)
add_para("Testing plan: unit tests for backend controllers, integration tests for API contracts (use Postman/newman or pytest + requests), end-to-end tests using Playwright or Cypress. Include test runs in CI pipelines and report results.", first_line_indent=True)
add_para("Maintenance: maintain Ansible inventories, set runbooks for incident response, document Nagios thresholds and escalation contacts.", first_line_indent=True)

doc.add_page_break()

# REFERENCES & APPENDICES
add_heading("REFERENCES", level=1)
add_para("References: React docs, Node.js docs, python-docx, Ansible docs, Nagios docs, articles on fuzzy search and string matching.", first_line_indent=True)

add_heading("APPENDIX A: FULL ANSIBLE PLAYBOOK (example)", level=1)
add_para("Append full playbook contents and usage notes here.", first_line_indent=True)

add_heading("APPENDIX B: NAGIOS CONFIG SNIPPETS", level=1)
add_para("Add nagios.cfg extracts, services.cfg entries, and custom check scripts. Include steps to install NRPE and add remote host configs.", first_line_indent=True)

add_heading("APPENDIX C: AI MODULE DETAILS", level=1)
add_para("Include model hyperparameters if you fine-tune a model, training dataset notes (synthetic misspellings), and evaluation metrics (precision/recall for suggestion correctness).", first_line_indent=True)

# Footer page numbering
add_page_number_footer()

# Save
doc.save(OUTPUT_FILENAME)
print(f"Saved: {OUTPUT_FILENAME}")
print("Open the document in MS Word and update the Table of Contents (right-click → Update Field or press F9).")
